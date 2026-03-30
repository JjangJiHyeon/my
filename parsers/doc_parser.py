"""
Parser for .doc (binary Word 97-2003) and .docx (Open XML) files.

Strategy order for .doc:
  1. Apache Tika  (most robust, needs Java)
  2. pywin32 COM  (Windows + Word installed)
  3. OLE binary piece-table  (pure-Python fallback)
"""

from __future__ import annotations

import logging
import os
import re
import struct
from typing import Any

logger = logging.getLogger(__name__)


# ── text normalisation (shared) ──────────────────────────────────────

def _normalise_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── public entry point ───────────────────────────────────────────────

def parse_doc(filepath: str) -> dict[str, Any]:
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".docx":
        return _parse_docx(filepath)

    errors: list[str] = []

    for name, fn in [
        ("tika", _parse_doc_tika),
        ("win32com", _parse_doc_win32),
        ("OLE binary", _parse_doc_binary),
    ]:
        try:
            result = fn(filepath)
            result.setdefault("metadata", {})["parse_strategies_tried"] = (
                errors if errors else [f"{name} (first attempt)"]
            )
            return result
        except Exception as exc:
            msg = f"{name}: {exc}"
            errors.append(msg)
            logger.info("doc strategy '%s' failed: %s", name, exc)

    return {
        "pages": [],
        "metadata": {"parse_strategies_tried": errors},
        "status": "error",
        "error": "All parsing strategies failed: " + " | ".join(errors),
    }


# ── Strategy: python-docx (.docx) ───────────────────────────────────

def _parse_docx(filepath: str) -> dict[str, Any]:
    from docx import Document

    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = _normalise_text("\n".join(paragraphs))

    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    metadata: dict[str, Any] = {
        "page_count": 1,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "char_count": len(full_text),
        "parser_used": "python-docx",
    }

    pages = [{"page_num": 1, "text": full_text, "tables": tables}]
    return {"pages": pages, "metadata": metadata, "status": "success"}


# ── Strategy 1: Tika ────────────────────────────────────────────────

def _parse_doc_tika(filepath: str) -> dict[str, Any]:
    try:
        from tika import parser as tika_parser
    except ImportError:
        raise RuntimeError("tika package not installed")

    parsed = tika_parser.from_file(filepath)
    content = parsed.get("content")
    if not content or not content.strip():
        raise ValueError("Tika returned empty content")

    full_text = _normalise_text(content)
    tika_meta = parsed.get("metadata") or {}

    metadata: dict[str, Any] = {
        "page_count": int(tika_meta.get("xmpTPg:NPages", 1) or 1),
        "char_count": len(full_text),
        "parser_used": "Apache Tika",
        "tika_content_type": tika_meta.get("Content-Type", ""),
    }

    pages = [{"page_num": 1, "text": full_text, "tables": []}]
    return {"pages": pages, "metadata": metadata, "status": "success"}


# ── Strategy 2: pywin32 COM ─────────────────────────────────────────

def _parse_doc_win32(filepath: str) -> dict[str, Any]:
    try:
        import win32com.client  # type: ignore
    except ImportError:
        raise RuntimeError(
            "pywin32 is not installed. Install with: pip install pywin32  "
            "(requires Microsoft Word on this machine)"
        )

    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(filepath), ReadOnly=True)
        full_text = _normalise_text(doc.Content.Text)
        doc.Close(False)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass

    metadata: dict[str, Any] = {
        "page_count": 1,
        "char_count": len(full_text),
        "parser_used": "win32com (Microsoft Word)",
    }
    pages = [{"page_num": 1, "text": full_text, "tables": []}]
    return {"pages": pages, "metadata": metadata, "status": "success"}


# ── Strategy 3: OLE binary piece-table ──────────────────────────────

def _parse_doc_binary(filepath: str) -> dict[str, Any]:
    import olefile

    if not olefile.isOleFile(filepath):
        raise ValueError("Not an OLE2 compound file")

    ole = olefile.OleFileIO(filepath)
    try:
        full_text = _extract_ole_text(ole)
    finally:
        ole.close()

    full_text = _normalise_text(full_text)

    metadata: dict[str, Any] = {
        "page_count": 1,
        "char_count": len(full_text),
        "parser_used": "OLE binary (piece table)",
    }
    pages = [{"page_num": 1, "text": full_text, "tables": []}]
    return {"pages": pages, "metadata": metadata, "status": "success"}


def _extract_ole_text(ole: Any) -> str:
    word_doc = ole.openstream("WordDocument").read()

    magic = struct.unpack_from("<H", word_doc, 0)[0]
    if magic not in (0xA5EC, 0xA5DC):
        raise ValueError(f"Invalid Word magic: {hex(magic)}")

    flags = struct.unpack_from("<H", word_doc, 0x000A)[0]
    table_name = "1Table" if (flags & 0x0200) else "0Table"

    if not ole.exists(table_name):
        alt = "0Table" if table_name == "1Table" else "1Table"
        if ole.exists(alt):
            table_name = alt
        else:
            raise ValueError("No Table stream found")

    table_stream = ole.openstream(table_name).read()

    csw = struct.unpack_from("<H", word_doc, 0x20)[0]
    fibrg_w_end = 0x22 + csw * 2

    cslw = struct.unpack_from("<H", word_doc, fibrg_w_end)[0]
    fibrg_lw_start = fibrg_w_end + 2
    fibrg_lw_end = fibrg_lw_start + cslw * 4

    ccpText = struct.unpack_from("<i", word_doc, fibrg_lw_start)[0]

    cb_rg = struct.unpack_from("<H", word_doc, fibrg_lw_end)[0]
    fclcb_start = fibrg_lw_end + 2

    if cb_rg < 34:
        raise ValueError(f"FIBRgFcLcb too small ({cb_rg} pairs)")

    fc_clx = struct.unpack_from("<I", word_doc, fclcb_start + 33 * 8)[0]
    lcb_clx = struct.unpack_from("<I", word_doc, fclcb_start + 33 * 8 + 4)[0]

    if lcb_clx == 0:
        raise ValueError("CLX has zero length")

    clx = table_stream[fc_clx: fc_clx + lcb_clx]

    off = 0
    while off < len(clx) and clx[off] == 0x01:
        cb = struct.unpack_from("<H", clx, off + 1)[0]
        off += 3 + cb

    if off >= len(clx) or clx[off] != 0x02:
        raise ValueError("Pcdt marker (0x02) not found")

    off += 1
    lcb_pcdt = struct.unpack_from("<I", clx, off)[0]
    off += 4
    plc = clx[off: off + lcb_pcdt]

    n_pieces = (lcb_pcdt - 4) // 12
    if n_pieces <= 0:
        raise ValueError("No text pieces found")

    cps = [struct.unpack_from("<I", plc, i * 4)[0] for i in range(n_pieces + 1)]
    pcd_base = (n_pieces + 1) * 4

    parts: list[str] = []
    for i in range(n_pieces):
        cp_start, cp_end = cps[i], cps[i + 1]
        if cp_start >= ccpText:
            break
        n_chars = min(cp_end, ccpText) - cp_start

        pcd = plc[pcd_base + i * 8: pcd_base + (i + 1) * 8]
        fc_raw = struct.unpack_from("<I", pcd, 2)[0]
        compressed = bool(fc_raw & (1 << 30))
        fc_val = fc_raw & 0x3FFFFFFF

        if compressed:
            boff = fc_val // 2
            parts.append(word_doc[boff: boff + n_chars].decode("cp1252", errors="replace"))
        else:
            boff = fc_val
            parts.append(word_doc[boff: boff + n_chars * 2].decode("utf-16-le", errors="replace"))

    return "".join(parts)
